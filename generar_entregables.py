"""
Generador de Documentos Entregables — OsterwalderAI Architect
Genera los archivos Word de los entregables académicos.

Archivos generados:
1. Informe_Proceso_Practica.docx    — Informe del proceso de elaboración
2. Procedimiento_Uso_GPT.docx       — Manual de uso detallado del sistema
3. Informe_Comparativo_IA.docx      — Plantilla del informe comparativo (se llena con resultados de la IA)
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

ENTREGABLES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "entregables")
os.makedirs(ENTREGABLES_DIR, exist_ok=True)

FECHA_ACTUAL = datetime.now().strftime("%d de %B de %Y")
CURSO = "Sistemas Analíticos"
UNIVERSIDAD = "Universidad Nacional de Ingeniería (UNI)"
CICLO = "2026-I"

# ══════════════════════════════════════════════════════════════
# UTILIDADES
# ══════════════════════════════════════════════════════════════

def configurar_estilos(doc):
    """Configura estilos base del documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

def agregar_portada(doc, titulo, subtitulo=""):
    """Agrega una portada profesional al documento."""
    # Espacio superior
    for _ in range(4):
        doc.add_paragraph()

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OsterwalderAI Architect")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x00, 0x96, 0xC7)
    run.font.bold = True

    # Subtítulo del sistema
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Consultor de Negocios Inteligente con IA")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Línea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = RGBColor(0x00, 0x96, 0xC7)
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Título del documento
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    if subtitulo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitulo)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Espacio
    for _ in range(4):
        doc.add_paragraph()

    # Información del curso
    info_lines = [
        f"Curso: {CURSO}",
        f"Institución: {UNIVERSIDAD}",
        f"Ciclo: {CICLO}",
        f"Fecha: {FECHA_ACTUAL}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

def agregar_titulo_seccion(doc, texto, nivel=1):
    """Agrega un título de sección con formato."""
    heading = doc.add_heading(texto, level=nivel)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def agregar_parrafo(doc, texto, bold=False, italic=False):
    """Agrega un párrafo con formato opcional."""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.bold = bold
    run.font.italic = italic
    return p

def agregar_lista(doc, items, estilo='List Bullet'):
    """Agrega una lista con viñetas."""
    for item in items:
        doc.add_paragraph(item, style=estilo)


# ══════════════════════════════════════════════════════════════
# DOCUMENTO 1: INFORME DEL PROCESO
# ══════════════════════════════════════════════════════════════

def generar_informe_proceso():
    """Genera el Informe del Proceso de Elaboración de la Práctica."""
    doc = Document()
    configurar_estilos(doc)
    agregar_portada(doc,
        "Informe del Proceso de Elaboración",
        "Práctica Calificada 02 — Sistemas Analíticos"
    )

    # ─── Índice ───
    agregar_titulo_seccion(doc, "Índice de Contenidos")
    indice = [
        "1. Introducción y Objetivo",
        "2. Stack Tecnológico Utilizado",
        "3. Arquitectura del Sistema",
        "4. Proceso de Desarrollo Paso a Paso",
        "5. Configuración del Entorno de Desarrollo",
        "6. Implementación del Backend (FastAPI + RAG)",
        "7. Implementación del Frontend (Interfaz Web Premium)",
        "8. Integración de Modelos de IA Locales",
        "9. Generación de los 5 Modelos de Negocio",
        "10. Sistema de Exportación a PDF",
        "11. Pruebas y Validación",
        "12. Dificultades Encontradas y Soluciones",
        "13. Conclusiones",
        "14. Referencias",
    ]
    for item in indice:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # ─── 1. Introducción ───
    agregar_titulo_seccion(doc, "1. Introducción y Objetivo")
    agregar_parrafo(doc,
        "El presente documento detalla el proceso completo de diseño, desarrollo e implementación "
        "del proyecto OsterwalderAI Architect, un Consultor de Negocios Inteligente que opera 100% "
        "en modo offline. Este sistema fue desarrollado como parte de la Práctica Calificada 02 del "
        "curso de Sistemas Analíticos de la Universidad Nacional de Ingeniería (UNI), ciclo 2026-I."
    )
    agregar_parrafo(doc,
        "El objetivo principal es crear una herramienta de inteligencia artificial capaz de analizar "
        "documentos de negocio (tesis de maestría, planes de negocio, etc.) y generar automáticamente "
        "cinco modelos estratégicos fundamentales basados en la metodología de Alexander Osterwalder: "
        "Business Model Canvas, Mapa de Empatía, Mapa de Persona/Cliente, Lienzo de Propuesta de "
        "Valor y Customer Journey Map."
    )

    # ─── 2. Stack Tecnológico ───
    agregar_titulo_seccion(doc, "2. Stack Tecnológico Utilizado")
    agregar_parrafo(doc, "El sistema utiliza las siguientes tecnologías principales:", bold=True)

    # Tabla de tecnologías
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Componente'
    hdr[1].text = 'Tecnología'
    hdr[2].text = 'Propósito'

    tech_data = [
        ("Lenguaje", "Python 3.12", "Lenguaje principal del proyecto"),
        ("Backend", "FastAPI + Uvicorn", "Servidor web de alto rendimiento"),
        ("IA / RAG", "LangChain v0.2+", "Orquestación del pipeline RAG"),
        ("Base Vectorial", "ChromaDB", "Almacenamiento de embeddings"),
        ("Documentos", "PyPDF", "Procesamiento de archivos PDF"),
        ("Modelo LLM", "GPT4All (Llama-3.2-1B-Instruct Q4)", "Inferencia local offline"),
        ("Embeddings", "all-MiniLM-L6-v2 (GGUF)", "Vectorización de texto"),
        ("Frontend", "HTML5 + CSS3 + JavaScript", "Interfaz web premium"),
        ("Diagramas", "Mermaid.js (local)", "Renderizado de modelos visuales"),
        ("Reportes", "fpdf2", "Generación de PDFs profesionales"),
        ("Desktop", "pywebview", "Wrapper de aplicación nativa"),
    ]
    for comp, tech, prop in tech_data:
        row = table.add_row().cells
        row[0].text = comp
        row[1].text = tech
        row[2].text = prop

    # ─── 3. Arquitectura del Sistema ───
    doc.add_page_break()
    agregar_titulo_seccion(doc, "3. Arquitectura del Sistema")
    agregar_parrafo(doc,
        "El sistema sigue una arquitectura cliente-servidor local con los siguientes componentes:"
    )
    agregar_parrafo(doc, "Flujo de datos del sistema:", bold=True)
    flujo = [
        "1. El usuario carga un documento PDF a través de la interfaz web.",
        "2. El backend procesa el PDF con PyPDF y lo fragmenta con RecursiveCharacterTextSplitter.",
        "3. Los fragmentos se vectorizan con el modelo all-MiniLM-L6-v2 (embeddings locales).",
        "4. Los vectores se almacenan en ChromaDB como base de conocimiento.",
        "5. Al solicitar un modelo de negocio, el sistema realiza una búsqueda semántica (RAG).",
        "6. El contexto recuperado se combina con prompts especializados por modelo.",
        "7. El LLM local (Llama-3.2-1B) genera la respuesta estructurada.",
        "8. El frontend parsea la respuesta y construye diagramas Mermaid.js.",
        "9. Los resultados pueden exportarse a PDF profesional con fpdf2.",
    ]
    for paso in flujo:
        doc.add_paragraph(paso, style='List Number')

    # ─── 4. Proceso de Desarrollo ───
    agregar_titulo_seccion(doc, "4. Proceso de Desarrollo Paso a Paso")
    agregar_parrafo(doc,
        "El desarrollo del proyecto se realizó de forma iterativa con asistencia de IA "
        "(Antigravity AI / Claude), siguiendo estos pasos principales:"
    )

    pasos_dev = [
        ("Paso 1: Definición de Requerimientos",
         "Se analizó el enunciado de la PC02 y se definieron los requerimientos funcionales: "
         "5 modelos de negocio, chat con IA, comparativa, exportación a PDF y funcionamiento 100% offline."),
        ("Paso 2: Creación de la Estructura del Proyecto",
         "Se creó la estructura de carpetas organizada: static/, knowledge_base/, uploads/, models/, "
         "chroma_db/, exports/, entregables/."),
        ("Paso 3: Configuración de Dependencias",
         "Se definió el requirements.txt con todas las librerías necesarias: FastAPI, LangChain, "
         "ChromaDB, GPT4All, fpdf2, pywebview, entre otras."),
        ("Paso 4: Desarrollo del Backend (app.py)",
         "Se implementó el servidor FastAPI con: descarga automática de modelos .gguf, pipeline RAG "
         "completo (LangChain + ChromaDB + GPT4All), endpoints especializados para cada modelo de negocio, "
         "y sistema de generación de PDF con fpdf2."),
        ("Paso 5: Desarrollo del Frontend",
         "Se creó la interfaz web premium con diseño Dark Mode, Glassmorphism, animaciones fluidas, "
         "zona drag & drop para PDFs, renderizado de diagramas Mermaid.js y botones de exportar a PDF."),
        ("Paso 6: Integración y Pruebas",
         "Se integró el frontend con el backend, se probó la carga de documentos, generación de modelos, "
         "exportación y se validó el funcionamiento offline completo."),
        ("Paso 7: Documentación y Entregables",
         "Se generaron los documentos Word de entregables: informe del proceso, procedimiento de uso "
         "e informe comparativo."),
    ]
    for titulo_paso, descripcion in pasos_dev:
        agregar_titulo_seccion(doc, titulo_paso, nivel=2)
        agregar_parrafo(doc, descripcion)

    # ─── 5-8. Secciones de implementación ───
    doc.add_page_break()
    agregar_titulo_seccion(doc, "5. Configuración del Entorno de Desarrollo")
    agregar_parrafo(doc,
        "Para ejecutar el proyecto, se requiere Python 3.12 en un sistema Windows. "
        "El proceso de configuración es el siguiente:"
    )
    agregar_lista(doc, [
        "Clonar el repositorio: git clone https://github.com/MartinZapanaBerrospi/osterwalderAI-architect.git",
        "Crear un entorno virtual: python -m venv venv",
        "Activar el entorno: venv\\Scripts\\activate",
        "Instalar dependencias: pip install -r requirements.txt",
        "Ejecutar la aplicación: python app.py",
    ])
    agregar_parrafo(doc,
        "En la primera ejecución, el sistema descargará automáticamente los modelos de IA "
        "(Llama-3.2-1B-Instruct y all-MiniLM-L6-v2) que ocupan aproximadamente 1.2 GB. "
        "Posteriormente, el sistema funciona 100% offline."
    )

    agregar_titulo_seccion(doc, "6. Implementación del Backend (FastAPI + RAG)")
    agregar_parrafo(doc,
        "El backend se implementó en app.py con FastAPI como framework web. "
        "El sistema RAG (Retrieval-Augmented Generation) sigue este flujo:"
    )
    agregar_lista(doc, [
        "Descarga automática de modelos .gguf si no existen localmente.",
        "Carga de embeddings con GPT4AllEmbeddings (all-MiniLM-L6-v2).",
        "Indexación de documentos PDF en ChromaDB con fragmentos de 1000 caracteres.",
        "Configuración del LLM local (Llama-3.2-1B) con temperatura 0.4.",
        "Prompts especializados en formato Llama 3.2 para cada modelo de negocio.",
    ])

    agregar_titulo_seccion(doc, "7. Implementación del Frontend")
    agregar_parrafo(doc,
        "La interfaz web se desarrolló con HTML5, CSS3 puro y JavaScript vanilla, "
        "siguiendo principios de diseño premium:"
    )
    agregar_lista(doc, [
        "Dark Mode con paleta de colores oscuros (#07070d, #0d0d18) y acentos cyan/púrpura.",
        "Glassmorphism mediante backdrop-filter: blur(24px) y bordes translúcidos.",
        "Animaciones fluidas con keyframes CSS (fadeInUp, scaleIn, shimmer, statusPulse).",
        "Navegación por sidebar con 5 vistas principales.",
        "Indicador de estado del sistema en tiempo real con barra de progreso.",
        "Zona de arrastrar y soltar (drag & drop) para carga de PDFs.",
        "Renderizado de diagramas Mermaid.js (mindmaps y flowcharts) para cada modelo de negocio.",
        "Sistema de notificaciones toast para feedback visual al usuario.",
    ])

    agregar_titulo_seccion(doc, "8. Integración de Modelos de IA Locales")
    agregar_parrafo(doc,
        "Se utilizan dos modelos de IA en formato .gguf que corren 100% localmente:"
    )
    agregar_lista(doc, [
        "Llama-3.2-1B-Instruct-Q4_0.gguf: Modelo de lenguaje para generación de texto. "
        "Cuantizado a 4 bits para eficiencia en hardware limitado.",
        "all-MiniLM-L6-v2.gguf2.f16.gguf: Modelo de embeddings para vectorización de texto. "
        "Utilizado para la búsqueda semántica en ChromaDB.",
    ])

    # ─── 9. Modelos de Negocio ───
    doc.add_page_break()
    agregar_titulo_seccion(doc, "9. Generación de los 5 Modelos de Negocio")
    modelos = [
        ("Business Model Canvas (BMC)",
         "9 bloques: Segmentos de Clientes, Propuesta de Valor, Canales, Relaciones con Clientes, "
         "Fuentes de Ingresos, Recursos Clave, Actividades Clave, Asociaciones Clave, Estructura de Costos."),
        ("Mapa de Empatía",
         "6 bloques: Piensa y Siente, Qué Ve, Qué Oye, Dice y Hace, Esfuerzos (Pains), Resultados (Gains)."),
        ("Mapa de Persona / Cliente",
         "6 bloques: Perfil, Demografía, Comportamiento, Necesidades, Frustraciones, Motivaciones."),
        ("Lienzo de Propuesta de Valor",
         "6 bloques divididos en Perfil del Cliente (Trabajos, Dolores, Ganancias) y "
         "Mapa de Valor (Productos/Servicios, Aliviadores de Dolor, Creadores de Ganancia)."),
        ("Customer Journey Map",
         "8 etapas en 3 fases: ANTES (Descubrimiento, Consideración, Decisión), "
         "DURANTE (Compra, Uso), DESPUÉS (Soporte, Retención, Recomendación)."),
    ]
    for nombre, descripcion in modelos:
        agregar_titulo_seccion(doc, nombre, nivel=2)
        agregar_parrafo(doc, descripcion)

    # ─── 10-14 ───
    agregar_titulo_seccion(doc, "10. Sistema de Exportación a PDF")
    agregar_parrafo(doc,
        "El sistema utiliza la librería fpdf2 para generar reportes PDF profesionales con: "
        "portada con título y fecha, secciones por cada modelo generado, formato limpio con "
        "tipografía Helvetica y el branding de OsterwalderAI Architect."
    )

    agregar_titulo_seccion(doc, "11. Pruebas y Validación")
    agregar_parrafo(doc,
        "Se realizaron pruebas funcionales de cada componente del sistema, verificando: "
        "la descarga automática de modelos, la indexación de PDFs, la generación correcta de "
        "los 5 modelos de negocio, la comparativa de factibilidad y la exportación a PDF."
    )

    agregar_titulo_seccion(doc, "12. Dificultades Encontradas y Soluciones")
    agregar_parrafo(doc,
        "Se documentan las dificultades técnicas encontradas durante el desarrollo:"
    )
    agregar_lista(doc, [
        "Rendimiento del LLM: El modelo de 1B parámetros tiene capacidad limitada. Se optimizaron "
        "los prompts para obtener respuestas estructuradas concisas.",
        "Funcionamiento offline: Se eliminaron todas las dependencias CDN (Google Fonts, Mermaid.js CDN) "
        "y se empaquetaron las librerías localmente.",
        "Compatibilidad de LangChain: Se ajustaron los imports para la versión 0.2+ de LangChain, "
        "que reorganizó los módulos en paquetes separados.",
    ])

    agregar_titulo_seccion(doc, "13. Conclusiones")
    agregar_parrafo(doc,
        "El proyecto OsterwalderAI Architect demuestra la viabilidad de crear herramientas "
        "de consultoría de negocios potenciadas por IA que operan completamente offline. "
        "El uso de modelos .gguf cuantizados permite ejecutar inferencia en hardware convencional, "
        "mientras que la combinación de RAG con ChromaDB enriquece las respuestas con contexto "
        "documental específico. La interfaz premium con Dark Mode y Glassmorphism proporciona "
        "una experiencia de usuario profesional y moderna."
    )

    agregar_titulo_seccion(doc, "14. Referencias")
    agregar_lista(doc, [
        "Osterwalder, A. & Pigneur, Y. (2010). Generación de Modelos de Negocio.",
        "Osterwalder, A. et al. (2014). Diseñando la Propuesta de Valor.",
        "LangChain Documentation: https://python.langchain.com/",
        "FastAPI Documentation: https://fastapi.tiangolo.com/",
        "GPT4All: https://gpt4all.io/",
        "ChromaDB: https://www.trychroma.com/",
    ])

    # Guardar
    filepath = os.path.join(ENTREGABLES_DIR, "Informe_Proceso_Practica.docx")
    doc.save(filepath)
    print(f"[OK] Generado: {filepath}")
    return filepath


# ══════════════════════════════════════════════════════════════
# DOCUMENTO 2: PROCEDIMIENTO DE USO DEL GPT
# ══════════════════════════════════════════════════════════════

def generar_procedimiento_uso():
    """Genera el Procedimiento Detallado para el Uso del Sistema."""
    doc = Document()
    configurar_estilos(doc)
    agregar_portada(doc,
        "Procedimiento Detallado para el Uso",
        "Manual de Usuario — OsterwalderAI Architect"
    )

    # ─── Índice ───
    agregar_titulo_seccion(doc, "Índice de Contenidos")
    indice = [
        "1. Descripción General del Sistema",
        "2. Requisitos del Sistema",
        "3. Instalación y Configuración",
        "4. Inicio de la Aplicación",
        "5. Interfaz del Usuario — Navegación",
        "6. Módulo: Consultar IA (Chat)",
        "7. Módulo: Analizar Documento",
        "8. Módulo: Comparativa de Factibilidad",
        "9. Módulo: Exportar a PDF",
        "10. Indicador de Estado del Sistema",
        "11. Solución de Problemas Frecuentes",
        "12. Glosario de Términos",
    ]
    for item in indice:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    # ─── 1. Descripción ───
    agregar_titulo_seccion(doc, "1. Descripción General del Sistema")
    agregar_parrafo(doc,
        "OsterwalderAI Architect es un Consultor de Negocios Inteligente que utiliza "
        "Inteligencia Artificial para analizar documentos de negocio y generar automáticamente "
        "cinco modelos estratégicos basados en la metodología de Alexander Osterwalder."
    )
    agregar_parrafo(doc, "El sistema genera los siguientes modelos:", bold=True)
    agregar_lista(doc, [
        "Business Model Canvas (9 bloques de Osterwalder)",
        "Mapa de Empatía (comprensión emocional del cliente)",
        "Mapa de Persona / Cliente (perfil del buyer persona)",
        "Lienzo de Propuesta de Valor (ajuste producto-mercado)",
        "Customer Journey Map (experiencia del cliente: Antes, Durante, Después)",
    ])
    agregar_parrafo(doc,
        "Características principales: Funciona 100% offline, no requiere conexión a internet "
        "después de la instalación inicial. Utiliza modelos de IA locales en formato .gguf."
    )

    # ─── 2. Requisitos ───
    agregar_titulo_seccion(doc, "2. Requisitos del Sistema")
    agregar_parrafo(doc, "Requisitos mínimos de hardware:", bold=True)
    agregar_lista(doc, [
        "Procesador: Intel Core i5 / AMD Ryzen 5 o superior",
        "Memoria RAM: 8 GB mínimo (16 GB recomendado)",
        "Almacenamiento: 3 GB disponibles (1.2 GB para modelos de IA + proyecto)",
        "Sistema Operativo: Windows 10/11 (64 bits)",
    ])
    agregar_parrafo(doc, "Requisitos de software:", bold=True)
    agregar_lista(doc, [
        "Python 3.12 instalado y configurado en el PATH del sistema",
        "Git (opcional, para clonar el repositorio)",
        "Conexión a internet solo para la instalación inicial (descarga de modelos)",
    ])

    # ─── 3. Instalación ───
    agregar_titulo_seccion(doc, "3. Instalación y Configuración")
    agregar_titulo_seccion(doc, "3.1 Clonar el Repositorio", nivel=2)
    agregar_parrafo(doc,
        "Abra una terminal (PowerShell o CMD) y ejecute:"
    )
    agregar_parrafo(doc,
        "git clone https://github.com/MartinZapanaBerrospi/osterwalderAI-architect.git",
        bold=True
    )
    agregar_parrafo(doc,
        "cd osterwalderAI-architect",
        bold=True
    )

    agregar_titulo_seccion(doc, "3.2 Crear Entorno Virtual", nivel=2)
    agregar_parrafo(doc, "python -m venv venv", bold=True)
    agregar_parrafo(doc, "venv\\Scripts\\activate", bold=True)

    agregar_titulo_seccion(doc, "3.3 Instalar Dependencias", nivel=2)
    agregar_parrafo(doc, "pip install -r requirements.txt", bold=True)
    agregar_parrafo(doc,
        "Este paso instalará todas las librerías necesarias: FastAPI, LangChain, ChromaDB, "
        "GPT4All, fpdf2, pywebview, entre otras."
    )

    agregar_titulo_seccion(doc, "3.4 Primera Ejecución (Descarga de Modelos)", nivel=2)
    agregar_parrafo(doc,
        "En la primera ejecución, el sistema descargará automáticamente dos modelos de IA:"
    )
    agregar_lista(doc, [
        "Llama-3.2-1B-Instruct-Q4_0.gguf (~800 MB) — Modelo de lenguaje",
        "all-MiniLM-L6-v2.gguf2.f16.gguf (~45 MB) — Modelo de embeddings",
    ])
    agregar_parrafo(doc,
        "IMPORTANTE: Se requiere conexión a internet SOLO para esta descarga inicial. "
        "Después de esto, el sistema funciona 100% offline.",
        bold=True
    )

    # ─── 4. Inicio ───
    doc.add_page_break()
    agregar_titulo_seccion(doc, "4. Inicio de la Aplicación")
    agregar_parrafo(doc, "Para iniciar la aplicación, ejecute:", bold=True)
    agregar_parrafo(doc, "python app.py", bold=True)
    agregar_parrafo(doc,
        "La aplicación abrirá automáticamente una ventana de escritorio con la interfaz web. "
        "También puede acceder desde cualquier navegador visitando: http://127.0.0.1:8000"
    )
    agregar_parrafo(doc,
        "El sistema iniciará un proceso de inicialización que incluye: verificación de modelos, "
        "carga de embeddings, indexación de la base de conocimiento y carga del modelo de lenguaje. "
        "Este proceso puede tardar entre 30 segundos y 2 minutos dependiendo del hardware."
    )

    # ─── 5. Navegación ───
    agregar_titulo_seccion(doc, "5. Interfaz del Usuario — Navegación")
    agregar_parrafo(doc,
        "La interfaz se compone de un sidebar de navegación a la izquierda y un área "
        "de contenido principal a la derecha. Las secciones disponibles son:"
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Icono'
    hdr[1].text = 'Sección'
    hdr[2].text = 'Función'

    nav_data = [
        ("🏠", "Inicio", "Panel principal con accesos rápidos a todas las funciones"),
        ("💬", "Consultar IA", "Chat libre con la IA sobre modelos de negocio"),
        ("🔬", "Analizar Documento", "Subir PDFs y generar los 5 modelos automáticamente"),
        ("⚖️", "Comparativa", "Comparar propuesta IA vs propuesta original de la tesis"),
        ("📄", "Exportar PDF", "Generar reportes profesionales en formato PDF"),
    ]
    for icono, seccion, funcion in nav_data:
        row = table.add_row().cells
        row[0].text = icono
        row[1].text = seccion
        row[2].text = funcion

    # ─── 6. Chat ───
    doc.add_page_break()
    agregar_titulo_seccion(doc, "6. Módulo: Consultar IA (Chat)")
    agregar_parrafo(doc, "Pasos para usar el chat:", bold=True)
    agregar_lista(doc, [
        "Haga clic en 'Consultar IA' en el sidebar.",
        "Espere a que el indicador de estado muestre 'Sistema Listo' (badge verde).",
        "Escriba su consulta en el campo de texto inferior.",
        "Presione Enter o haga clic en 'Enviar'.",
        "La IA procesará su consulta usando la base de conocimiento (RAG) y responderá.",
        "Cada respuesta tiene un botón 'Exportar a PDF' para guardar individualmente.",
    ])
    agregar_parrafo(doc, "Ejemplos de consultas:", bold=True)
    agregar_lista(doc, [
        "'¿Cuáles son los 9 bloques del Business Model Canvas?'",
        "'Explica la diferencia entre propuesta de valor y segmento de clientes.'",
        "'¿Cómo diseñar un Customer Journey Map para una fintech?'",
        "'Resume los conceptos clave del libro Generación de Modelos de Negocio.'",
    ])

    # ─── 7. Analizar Documento ───
    agregar_titulo_seccion(doc, "7. Módulo: Analizar Documento")
    agregar_parrafo(doc, "Este es el módulo principal del sistema. Pasos:", bold=True)
    pasos_analisis = [
        "Paso 1: Haga clic en 'Analizar Documento' en el sidebar.",
        "Paso 2: Arrastre un archivo PDF a la zona de upload, o haga clic para seleccionar.",
        "Paso 3: Espere a que el sistema procese e indexe el documento (se mostrará un mensaje de confirmación).",
        "Paso 4: Haga clic en el botón '🚀 Generar los 5 Modelos de Negocio'.",
        "Paso 5: El sistema generará cada modelo secuencialmente (mostrando progreso).",
        "Paso 6: Al completarse, verá cada modelo con sus bloques de información y diagrama Mermaid.",
        "Paso 7: Use los botones 'PDF' en cada tarjeta para exportar modelos individuales.",
    ]
    for paso in pasos_analisis:
        doc.add_paragraph(paso, style='List Number')

    agregar_parrafo(doc,
        "NOTA: La generación de cada modelo puede tardar entre 30 segundos y 3 minutos "
        "dependiendo del hardware. No cierre la aplicación durante el proceso.",
        italic=True
    )

    # ─── 8. Comparativa ───
    agregar_titulo_seccion(doc, "8. Módulo: Comparativa de Factibilidad")
    agregar_parrafo(doc, "Pasos para la comparación:", bold=True)
    agregar_lista(doc, [
        "Primero, genere los 5 modelos en el módulo 'Analizar Documento'.",
        "Vaya a la sección 'Comparativa' en el sidebar.",
        "En el campo de texto, pegue un resumen de la propuesta original de la tesis.",
        "Haga clic en '⚖️ Ejecutar Comparación de Factibilidad'.",
        "El sistema comparará ambas propuestas y emitirá un juicio que incluye:",
    ])
    agregar_lista(doc, [
        "Coincidencias entre ambas propuestas",
        "Diferencias identificadas",
        "Factibilidad Técnica (Alta/Media/Baja + justificación)",
        "Factibilidad Comercial (Alta/Media/Baja + justificación)",
        "Recomendaciones de mejora",
    ])

    # ─── 9. Exportar ───
    agregar_titulo_seccion(doc, "9. Módulo: Exportar a PDF")
    agregar_parrafo(doc, "Pasos para exportar:", bold=True)
    agregar_lista(doc, [
        "Vaya a 'Exportar PDF' en el sidebar.",
        "Escriba el título del reporte (por defecto: 'Análisis de Modelo de Negocio').",
        "Seleccione las secciones a incluir marcando/desmarcando checkboxes.",
        "Haga clic en '📄 Descargar Reporte PDF'.",
        "El archivo PDF se descargará automáticamente a su carpeta de descargas.",
    ])

    # ─── 10. Estado ───
    agregar_titulo_seccion(doc, "10. Indicador de Estado del Sistema")
    agregar_parrafo(doc,
        "El sistema muestra su estado actual en el sidebar con un indicador de color y barra de progreso:"
    )
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Light Grid Accent 1'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Color'
    hdr2[1].text = 'Estado'
    hdr2[2].text = 'Significado'
    estados = [
        ("🟡 Amarillo", "Descargando modelos", "Primera ejecución — descargando modelos de IA"),
        ("🔵 Azul", "Inicializando", "Cargando embeddings, indexando o cargando LLM"),
        ("🟢 Verde", "Sistema Listo", "Todos los módulos operativos — listo para usar"),
        ("🔴 Rojo", "Error", "Ha ocurrido un error — revise la consola"),
    ]
    for color, estado, significado in estados:
        row = table2.add_row().cells
        row[0].text = color
        row[1].text = estado
        row[2].text = significado

    # ─── 11. Solución de Problemas ───
    doc.add_page_break()
    agregar_titulo_seccion(doc, "11. Solución de Problemas Frecuentes")

    problemas = [
        ("El sistema se queda en 'Descargando modelos'",
         "Verifique su conexión a internet. Los modelos se descargan solo la primera vez. "
         "Si la descarga se interrumpió, elimine la carpeta 'models/' y reinicie."),
        ("Error: 'El sistema aún se está inicializando'",
         "Espere a que el indicador cambie a verde (Sistema Listo). La inicialización puede "
         "tardar hasta 2 minutos en equipos con hardware limitado."),
        ("Los diagramas Mermaid no se muestran",
         "Verifique que el archivo static/vendor/mermaid.min.js existe. Si no, copie la librería "
         "desde la distribución del proyecto."),
        ("La generación de modelos es muy lenta",
         "Es normal que cada modelo tarde 1-3 minutos con un CPU estándar. El modelo de 1B "
         "parámetros está optimizado para hardware limitado."),
        ("Error al exportar PDF",
         "Verifique que la carpeta 'exports/' existe y tiene permisos de escritura."),
    ]
    for prob, sol in problemas:
        agregar_titulo_seccion(doc, prob, nivel=2)
        agregar_parrafo(doc, f"Solución: {sol}")

    # ─── 12. Glosario ───
    agregar_titulo_seccion(doc, "12. Glosario de Términos")
    glosario = [
        ("RAG", "Retrieval-Augmented Generation — Técnica de IA que combina búsqueda de información con generación de texto."),
        ("LLM", "Large Language Model — Modelo de lenguaje grande capaz de comprender y generar texto."),
        ("GGUF", "Formato de archivo para modelos de IA cuantizados que permite ejecución local eficiente."),
        ("Embeddings", "Representaciones vectoriales de texto que permiten búsqueda semántica."),
        ("ChromaDB", "Base de datos vectorial ligera para almacenar e indexar embeddings."),
        ("BMC", "Business Model Canvas — Herramienta de diseño de modelos de negocio de Osterwalder."),
        ("Glassmorphism", "Estilo de diseño UI con efectos de vidrio esmerilado y transparencias."),
    ]
    for termino, definicion in glosario:
        agregar_titulo_seccion(doc, termino, nivel=2)
        agregar_parrafo(doc, definicion)

    # Guardar
    filepath = os.path.join(ENTREGABLES_DIR, "Procedimiento_Uso_GPT.docx")
    doc.save(filepath)
    print(f"[OK] Generado: {filepath}")
    return filepath


# ══════════════════════════════════════════════════════════════
# DOCUMENTO 3: INFORME COMPARATIVO (Plantilla)
# ══════════════════════════════════════════════════════════════

def generar_informe_comparativo():
    """Genera la plantilla del Informe Comparativo (se llena con resultados de la IA)."""
    doc = Document()
    configurar_estilos(doc)
    agregar_portada(doc,
        "Informe Comparativo Generado por IA",
        "Análisis de Factibilidad — Modelo de Negocio"
    )

    agregar_titulo_seccion(doc, "1. Resumen Ejecutivo")
    agregar_parrafo(doc,
        "[Este informe será completado con los resultados generados por OsterwalderAI Architect "
        "al analizar el documento de negocio y ejecutar la comparación de factibilidad.]"
    )

    agregar_titulo_seccion(doc, "2. Documento Analizado")
    agregar_parrafo(doc, "Nombre del documento: [A completar]")
    agregar_parrafo(doc, "Tipo: Tesis de maestría / Plan de negocio")
    agregar_parrafo(doc, "Páginas procesadas: [A completar]")
    agregar_parrafo(doc, "Fragmentos indexados: [A completar]")

    agregar_titulo_seccion(doc, "3. Modelos de Negocio Generados por IA")
    modelos = [
        "3.1 Business Model Canvas",
        "3.2 Mapa de Empatía",
        "3.3 Mapa de Persona / Cliente",
        "3.4 Lienzo de Propuesta de Valor",
        "3.5 Customer Journey Map",
    ]
    for modelo in modelos:
        agregar_titulo_seccion(doc, modelo, nivel=2)
        agregar_parrafo(doc, "[Contenido generado por la IA — pegar resultados aquí]")
        doc.add_paragraph()

    agregar_titulo_seccion(doc, "4. Análisis Comparativo")
    agregar_titulo_seccion(doc, "4.1 Coincidencias", nivel=2)
    agregar_parrafo(doc, "[A completar con resultados de la comparación de la IA]")
    agregar_titulo_seccion(doc, "4.2 Diferencias", nivel=2)
    agregar_parrafo(doc, "[A completar con resultados de la comparación de la IA]")
    agregar_titulo_seccion(doc, "4.3 Factibilidad Técnica", nivel=2)
    agregar_parrafo(doc, "[A completar: Alta/Media/Baja + justificación]")
    agregar_titulo_seccion(doc, "4.4 Factibilidad Comercial", nivel=2)
    agregar_parrafo(doc, "[A completar: Alta/Media/Baja + justificación]")

    agregar_titulo_seccion(doc, "5. Recomendaciones de la IA")
    agregar_parrafo(doc, "[A completar con las recomendaciones generadas por el sistema]")

    agregar_titulo_seccion(doc, "6. Conclusiones")
    agregar_parrafo(doc, "[A completar tras el análisis]")

    # Guardar
    filepath = os.path.join(ENTREGABLES_DIR, "Informe_Comparativo_IA.docx")
    doc.save(filepath)
    print(f"[OK] Generado: {filepath}")
    return filepath


# ══════════════════════════════════════════════════════════════
# EJECUTAR
# ══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 55)
    print("  Generando Documentos Entregables...")
    print("=" * 55)
    generar_informe_proceso()
    generar_procedimiento_uso()
    generar_informe_comparativo()
    print("=" * 55)
    print("  ¡Todos los entregables generados en entregables/!")
    print("=" * 55)
